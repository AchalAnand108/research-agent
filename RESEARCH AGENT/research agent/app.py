"""
╔══════════════════════════════════════════════════════════════════════════════╗
║              AI-POWERED RESEARCH AGENT — IBM watsonx.ai + Granite            ║
║                        Flask Backend  |  app.py                              ║
╚══════════════════════════════════════════════════════════════════════════════╝

AGENT_INSTRUCTIONS
==================
Customize the agent behavior below. All fields in this section directly control
how the AI model responds to every research request.

RESEARCH_DOMAIN:
    The primary domain / discipline the agent specializes in.
    Examples: "Computer Science & AI", "Biomedical Research",
              "Climate Science", "Economics", "General Academic Research"

WRITING_STYLE:
    Tone and voice the agent uses in responses.
    Options: "academic", "technical", "plain-language", "concise"

CITATION_FORMAT:
    Default citation style for references.
    Options: "APA", "IEEE", "MLA", "Chicago", "Harvard"

SOURCE_PREFERENCES:
    List of preferred source types, in order of priority.
    Trusted types: "peer-reviewed journals", "conference papers",
                   "preprints (arXiv)", "government reports", "textbooks"

SAFETY_RULES:
    Controls hallucination guard and uncertainty disclosure.
    - always_cite: Every factual claim must have a citation.
    - admit_uncertainty: Flag claims with low evidence as [UNCERTAIN].
    - no_fabricated_refs: Never invent paper titles/authors/DOIs.
    - trusted_sources_only: Prefer peer-reviewed and indexed sources.

MAX_PAPERS_PER_SEARCH:
    Maximum number of papers retrieved per arXiv/Scholar query.
    Recommended: 5–20.

REPORT_MAX_TOKENS:
    Maximum token budget for generated reports.
    Recommended: 1024–4096.
"""

# ─── AGENT INSTRUCTIONS (edit freely) ───────────────────────────────────────
AGENT_INSTRUCTIONS = {
    "RESEARCH_DOMAIN":      "General Academic Research",
    "WRITING_STYLE":        "academic",          # academic | technical | plain-language | concise
    "CITATION_FORMAT":      "APA",               # APA | IEEE | MLA | Chicago | Harvard
    "SOURCE_PREFERENCES":   [
        "peer-reviewed journals",
        "conference papers",
        "preprints (arXiv)",
        "government reports",
        "textbooks",
    ],
    "SAFETY_RULES": {
        "always_cite":          True,
        "admit_uncertainty":    True,
        "no_fabricated_refs":   True,
        "trusted_sources_only": True,
    },
    "MAX_PAPERS_PER_SEARCH": 10,
    "REPORT_MAX_TOKENS":     2048,
}
# ─────────────────────────────────────────────────────────────────────────────

import os
import json
import re
import uuid
import logging
import warnings
import io
from datetime import datetime
from pathlib import Path

from flask import (
    Flask, request, jsonify, render_template,
    send_from_directory, send_file, abort,
)
from flask_cors import CORS
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

# Silence non-fatal IBM SDK deprecation notices (api_deprecation, param_deprecation, disclaimer_warning)
warnings.filterwarnings("ignore", category=UserWarning, module="ibm_watsonx_ai")
warnings.filterwarnings("ignore", message=".*deprecated.*", category=Warning)

# ── IBM watsonx.ai ──────────────────────────────────────────────────────────
from ibm_watsonx_ai import APIClient, Credentials
from ibm_watsonx_ai.foundation_models import ModelInference
# metanames path is the same in 1.5.x; guard for any future rename
try:
    from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
except ImportError:                                   # pragma: no cover
    from ibm_watsonx_ai.foundation_models.utils.enums import GenTextParamsMetaNames as GenParams

# ── Document parsing ─────────────────────────────────────────────────────────
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import arxiv
    ARXIV_SUPPORT = True
except ImportError:
    ARXIV_SUPPORT = False

# ── Report generation ────────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        ListFlowable, ListItem,
    )
    from reportlab.lib import colors
    REPORTLAB_SUPPORT = True
except ImportError:
    REPORTLAB_SUPPORT = False

# ─────────────────────────────────────────────────────────────────────────────
# App setup
# ─────────────────────────────────────────────────────────────────────────────
# override=True ensures .env values always win over any stale shell env vars
load_dotenv(override=True)
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

app = Flask(__name__, template_folder="templates", static_folder="static")
CORS(app)

app.secret_key = os.getenv("FLASK_SECRET_KEY", uuid.uuid4().hex)
app.config["MAX_CONTENT_LENGTH"] = int(os.getenv("MAX_CONTENT_LENGTH", 16 * 1024 * 1024))

UPLOAD_FOLDER = Path(os.getenv("UPLOAD_FOLDER", "uploads"))
UPLOAD_FOLDER.mkdir(exist_ok=True)
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}

# ─────────────────────────────────────────────────────────────────────────────
# IBM watsonx.ai client
# ─────────────────────────────────────────────────────────────────────────────
# Read credentials via getenv *after* load_dotenv so they are always populated
IBM_API_KEY     = os.getenv("IBM_API_KEY", "")
IBM_PROJECT_ID  = os.getenv("IBM_PROJECT_ID", "")
IBM_WATSONX_URL = os.getenv("IBM_WATSONX_URL", "https://us-south.ml.cloud.ibm.com")

_watsonx_client: APIClient | None = None
_watsonx_model:  ModelInference | None = None


# Preferred instruct models in priority order (first one available wins).
# Override any of these by setting WATSONX_MODEL_ID in your .env file.
# Only IBM Granite models are used.
_PREFERRED_MODELS = [
    "ibm/granite-3-8b-instruct",       # us-south, eu-de, jp-tok (default)
    "ibm/granite-3-1-8b-instruct",     # newer alias
    "ibm/granite-3-2b-instruct",       # lightweight fallback
    "ibm/granite-13b-instruct-v2",     # larger, higher quality fallback
    "ibm/granite-8b-code-instruct",    # code-focused fallback
]

# Read optional override from .env
_MODEL_ID_OVERRIDE = os.getenv("WATSONX_MODEL_ID", "")

# Resolved model id (set on first successful init)
ACTIVE_MODEL_ID: str = _MODEL_ID_OVERRIDE or _PREFERRED_MODELS[0]


def get_watsonx_model() -> ModelInference | None:
    """Lazy-initialize the IBM watsonx.ai model.

    Tries each model in _PREFERRED_MODELS until one is accepted by the
    regional endpoint, so the app works across all IBM Cloud regions
    without any manual model-ID changes.
    """
    global _watsonx_client, _watsonx_model, ACTIVE_MODEL_ID
    if _watsonx_model:
        return _watsonx_model
    if not IBM_API_KEY or not IBM_PROJECT_ID:
        log.warning("IBM credentials not set — AI features will be disabled.")
        return None
    try:
        creds = Credentials(url=IBM_WATSONX_URL, api_key=IBM_API_KEY)
        _watsonx_client = APIClient(credentials=creds, project_id=IBM_PROJECT_ID)

        # If the user pinned a model, try only that one
        candidates = [_MODEL_ID_OVERRIDE] if _MODEL_ID_OVERRIDE else _PREFERRED_MODELS

        last_err = None
        for model_id in candidates:
            try:
                _watsonx_model = ModelInference(
                    model_id=model_id,
                    api_client=_watsonx_client,
                    project_id=IBM_PROJECT_ID,
                    params={
                        GenParams.DECODING_METHOD:    "greedy",
                        GenParams.MAX_NEW_TOKENS:     AGENT_INSTRUCTIONS["REPORT_MAX_TOKENS"],
                        GenParams.MIN_NEW_TOKENS:     10,
                        GenParams.TEMPERATURE:        0.3,
                        GenParams.TOP_P:              0.9,
                        GenParams.REPETITION_PENALTY: 1.1,
                        GenParams.STOP_SEQUENCES:     ["<|endoftext|>", "Human:", "User:"],
                    },
                )
                ACTIVE_MODEL_ID = model_id
                log.info("IBM watsonx.ai model initialized — %s", model_id)
                return _watsonx_model
            except Exception as e:
                log.warning("Model %s not available: %s", model_id, e)
                last_err = e
                _watsonx_model = None
                continue

        log.error("No supported model found for this region. Last error: %s", last_err)
        return None
    except Exception as exc:
        log.error("Failed to initialize watsonx client: %s", exc)
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Prompt builder
# ─────────────────────────────────────────────────────────────────────────────
def _system_prompt() -> str:
    ai = AGENT_INSTRUCTIONS
    safety = ai["SAFETY_RULES"]
    source_list = ", ".join(ai["SOURCE_PREFERENCES"])
    safety_text = (
        f"{'Always cite every factual claim. ' if safety['always_cite'] else ''}"
        f"{'Clearly mark claims with insufficient evidence as [UNCERTAIN]. ' if safety['admit_uncertainty'] else ''}"
        f"{'Never fabricate paper titles, authors, or DOIs. ' if safety['no_fabricated_refs'] else ''}"
        f"{'Prefer peer-reviewed and indexed sources only. ' if safety['trusted_sources_only'] else ''}"
    )
    return (
        f"You are an expert AI Research Agent specializing in {ai['RESEARCH_DOMAIN']}. "
        f"You assist researchers, academics, and students with literature reviews, paper summaries, "
        f"citation management, research gap identification, hypothesis generation, and report writing. "
        f"\n\nBehavior rules:"
        f"\n- Writing style: {ai['WRITING_STYLE']}."
        f"\n- Default citation format: {ai['CITATION_FORMAT']}."
        f"\n- Preferred sources (in order): {source_list}."
        f"\n- Safety: {safety_text}"
        f"\n- Structure all outputs with clear headings and bullet points where appropriate."
        f"\n- When listing references use the {ai['CITATION_FORMAT']} citation format exactly."
        f"\n- Identify research gaps and suggest future work when synthesizing literature."
        f"\n- Extract: authors, methodology, datasets, results, and limitations from papers."
    )


def call_granite(prompt: str, extra_context: str = "") -> str:
    """Send a prompt to IBM Granite and return the generated text."""
    model = get_watsonx_model()
    if not model:
        return (
            "⚠️ IBM watsonx.ai is not configured. "
            "Please add your IBM_API_KEY and IBM_PROJECT_ID to the .env file."
        )
    full_prompt = f"{_system_prompt()}\n\n"
    if extra_context:
        full_prompt += f"Context:\n{extra_context}\n\n"
    full_prompt += f"User: {prompt}\nAssistant:"
    try:
        response = model.generate_text(prompt=full_prompt)
        return response.strip() if response else "No response generated."
    except Exception as exc:
        log.error("Granite generation error: %s", exc)
        return f"Error communicating with IBM watsonx.ai: {exc}"


# ─────────────────────────────────────────────────────────────────────────────
# arXiv search helper
# ─────────────────────────────────────────────────────────────────────────────
def search_arxiv(query: str, max_results: int | None = None) -> list[dict]:
    """Search arXiv and return structured paper metadata."""
    if not ARXIV_SUPPORT:
        return []
    limit = max_results or AGENT_INSTRUCTIONS["MAX_PAPERS_PER_SEARCH"]
    try:
        client = arxiv.Client()
        search = arxiv.Search(query=query, max_results=limit,
                              sort_by=arxiv.SortCriterion.Relevance)
        papers = []
        for result in client.results(search):
            papers.append({
                "id":        result.entry_id,
                "title":     result.title,
                "authors":   [str(a) for a in result.authors[:5]],
                "abstract":  result.summary[:600] + ("…" if len(result.summary) > 600 else ""),
                "published": result.published.strftime("%Y-%m-%d") if result.published else "N/A",
                "url":       result.pdf_url or result.entry_id,
                "source":    "arXiv",
                "categories": result.categories[:3],
            })
        return papers
    except Exception as exc:
        log.warning("arXiv search error: %s", exc)
        return []


# ─────────────────────────────────────────────────────────────────────────────
# Document text extraction
# ─────────────────────────────────────────────────────────────────────────────
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_file(filepath: Path) -> str:
    suffix = filepath.suffix.lower()
    if suffix == ".pdf" and PDF_SUPPORT:
        text = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages[:30]:          # cap at 30 pages
                text.append(page.extract_text() or "")
        return "\n".join(text)[:8000]               # cap context length
    if suffix == ".docx" and DOCX_SUPPORT:
        doc = DocxDocument(str(filepath))
        return "\n".join(p.text for p in doc.paragraphs)[:8000]
    if suffix == ".txt":
        return filepath.read_text(encoding="utf-8", errors="ignore")[:8000]
    return ""


# ─────────────────────────────────────────────────────────────────────────────
# PDF report generation via ReportLab
# ─────────────────────────────────────────────────────────────────────────────
def _build_pdf_report(title: str, content: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle("Heading1Custom", parent=styles["Heading1"],
                                   textColor=colors.HexColor("#1a3a6b"), spaceAfter=8)
    body_style = ParagraphStyle("BodyCustom", parent=styles["BodyText"],
                                leading=16, spaceAfter=6)
    elements = []
    elements.append(Paragraph(title, heading_style))
    elements.append(Spacer(1, 0.4*cm))
    elements.append(HRFlowable(width="100%", thickness=1,
                                color=colors.HexColor("#3b82d4")))
    elements.append(Spacer(1, 0.4*cm))
    ts = datetime.now().strftime("%B %d, %Y %H:%M")
    elements.append(Paragraph(
        f"<font size='9' color='#57606a'>Generated by AI Research Agent · {ts}</font>",
        styles["Normal"]))
    elements.append(Spacer(1, 0.6*cm))
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 0.2*cm))
        elif line.startswith("## "):
            elements.append(Paragraph(line[3:], heading_style))
        elif line.startswith("# "):
            elements.append(Paragraph(line[2:], heading_style))
        elif line.startswith("- ") or line.startswith("* "):
            elements.append(Paragraph(f"• {line[2:]}", body_style))
        else:
            safe = re.sub(r"[<>&]", lambda m: {"<":"&lt;",">":"&gt;","&":"&amp;"}[m.group()], line)
            elements.append(Paragraph(safe, body_style))
    doc.build(elements)
    buf.seek(0)
    return buf.read()


def _build_docx_report(title: str, content: str) -> bytes:
    if not DOCX_SUPPORT:
        raise RuntimeError("python-docx not installed")
    doc = DocxDocument()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Generated by AI Research Agent · {datetime.now().strftime('%B %d, %Y %H:%M')}")
    doc.add_paragraph("")
    for line in content.split("\n"):
        line = line.strip()
        if line.startswith("## ") or line.startswith("# "):
            doc.add_heading(line.lstrip("# "), 2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# In-memory session storage (replace with DB for production)
# ─────────────────────────────────────────────────────────────────────────────
research_sessions: dict[str, dict] = {}   # session_id → {history, citations, title}


def _get_or_create_session(session_id: str) -> dict:
    if session_id not in research_sessions:
        research_sessions[session_id] = {
            "id":        session_id,
            "title":     "New Research Session",
            "history":   [],
            "citations": [],
            "created":   datetime.now().isoformat(),
        }
    return research_sessions[session_id]


# ═════════════════════════════════════════════════════════════════════════════
# ROUTES
# ═════════════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("index.html")


# ── Health check ──────────────────────────────────────────────────────────────
@app.route("/api/health")
def health():
    model_ok = get_watsonx_model() is not None
    return jsonify({
        "status":       "ok",
        "model":        ACTIVE_MODEL_ID,
        "model_ready":  model_ok,
        "arxiv":        ARXIV_SUPPORT,
        "pdf_support":  PDF_SUPPORT,
        "docx_support": DOCX_SUPPORT,
        "domain":       AGENT_INSTRUCTIONS["RESEARCH_DOMAIN"],
        "citation_fmt": AGENT_INSTRUCTIONS["CITATION_FORMAT"],
        "timestamp":    datetime.now().isoformat(),
    })


# ── Chat / Research Q&A ───────────────────────────────────────────────────────
@app.route("/api/chat", methods=["POST"])
def chat():
    data       = request.get_json(force=True)
    user_msg   = data.get("message", "").strip()
    session_id = data.get("session_id", str(uuid.uuid4()))
    if not user_msg:
        return jsonify({"error": "Empty message"}), 400

    session = _get_or_create_session(session_id)
    history = session["history"]

    # Build conversational context from last 6 exchanges
    ctx_lines = []
    for turn in history[-6:]:
        ctx_lines.append(f"User: {turn['user']}")
        ctx_lines.append(f"Assistant: {turn['assistant']}")
    context = "\n".join(ctx_lines)

    response = call_granite(user_msg, extra_context=context if context else "")

    history.append({"user": user_msg, "assistant": response,
                    "ts": datetime.now().isoformat()})
    if len(history) == 1:
        session["title"] = user_msg[:60]

    return jsonify({
        "response":   response,
        "session_id": session_id,
        "session_title": session["title"],
    })


# ── Paper Search (arXiv) ──────────────────────────────────────────────────────
@app.route("/api/search", methods=["POST"])
def search_papers():
    data  = request.get_json(force=True)
    query = data.get("query", "").strip()
    limit = min(int(data.get("limit", AGENT_INSTRUCTIONS["MAX_PAPERS_PER_SEARCH"])), 20)
    if not query:
        return jsonify({"error": "Query required"}), 400

    papers = search_arxiv(query, max_results=limit)

    # Generate an AI insight for the search results
    if papers:
        titles_block = "\n".join(
            f"[{i+1}] {p['title']} ({p['published']}) – {', '.join(p['authors'][:2])}"
            for i, p in enumerate(papers)
        )
        insight = call_granite(
            f"Given the following {len(papers)} papers retrieved for the query '{query}', "
            f"provide a brief (3-5 sentence) synthesis: identify the main research themes, "
            f"notable trends, and any apparent research gaps.",
            extra_context=titles_block,
        )
    else:
        insight = "No papers found. Try broader search terms."

    return jsonify({"papers": papers, "insight": insight, "query": query})


# ── Paper Summary ─────────────────────────────────────────────────────────────
@app.route("/api/summarize", methods=["POST"])
def summarize_paper():
    data  = request.get_json(force=True)
    title = data.get("title", "")
    abstract = data.get("abstract", "")
    content  = data.get("content", "")   # optional full text

    if not (title or abstract or content):
        return jsonify({"error": "Paper title, abstract, or content required"}), 400

    paper_text = f"Title: {title}\n\nAbstract: {abstract}\n\nContent excerpt: {content}"
    summary = call_granite(
        "Summarize this research paper. Include: (1) Research question / objective, "
        "(2) Methodology, (3) Key datasets used, (4) Main results and findings, "
        "(5) Limitations, (6) Contributions to the field, (7) Suggested citations.",
        extra_context=paper_text,
    )
    return jsonify({"summary": summary})


# ── Literature Review ─────────────────────────────────────────────────────────
@app.route("/api/literature-review", methods=["POST"])
def literature_review():
    data    = request.get_json(force=True)
    topic   = data.get("topic", "").strip()
    papers  = data.get("papers", [])          # list of {title, authors, abstract, year}
    session_id = data.get("session_id", str(uuid.uuid4()))
    if not topic:
        return jsonify({"error": "Topic required"}), 400

    # Auto-fetch papers if none provided
    if not papers and ARXIV_SUPPORT:
        papers = search_arxiv(topic)

    papers_block = "\n\n".join(
        f"Paper {i+1}:\n"
        f"  Title: {p.get('title','')}\n"
        f"  Authors: {', '.join(p.get('authors', [])[:3])}\n"
        f"  Year: {p.get('published','N/A')[:4]}\n"
        f"  Abstract: {p.get('abstract','')[:400]}"
        for i, p in enumerate(papers[:10])
    )

    review = call_granite(
        f"Write a structured literature review on the topic: '{topic}'. "
        f"Use the papers provided as primary sources. Include: "
        f"(1) Introduction and scope, (2) Thematic analysis of the literature, "
        f"(3) Comparison of methodologies, (4) Key findings synthesis, "
        f"(5) Identified research gaps, (6) Future research directions, "
        f"(7) Conclusion, (8) References in {AGENT_INSTRUCTIONS['CITATION_FORMAT']} format.",
        extra_context=papers_block,
    )

    session = _get_or_create_session(session_id)
    session["history"].append({
        "user":      f"Literature review: {topic}",
        "assistant": review,
        "ts":        datetime.now().isoformat(),
        "type":      "literature_review",
    })

    # Extract citations
    for p in papers:
        authors = p.get("authors", [])
        year    = p.get("published", "")[:4]
        title_p = p.get("title", "")
        url     = p.get("url", "")
        if title_p:
            session["citations"].append({
                "id":      str(uuid.uuid4()),
                "title":   title_p,
                "authors": authors,
                "year":    year,
                "url":     url,
                "source":  p.get("source", "arXiv"),
                "format":  AGENT_INSTRUCTIONS["CITATION_FORMAT"],
            })

    return jsonify({
        "review":     review,
        "papers":     papers,
        "session_id": session_id,
    })


# ── Compare Papers ────────────────────────────────────────────────────────────
@app.route("/api/compare", methods=["POST"])
def compare_papers():
    data   = request.get_json(force=True)
    papers = data.get("papers", [])    # list of {title, abstract, authors, year}
    if len(papers) < 2:
        return jsonify({"error": "At least two papers required for comparison"}), 400

    papers_block = "\n\n".join(
        f"Paper {chr(65+i)} — {p.get('title','Untitled')} "
        f"({', '.join(p.get('authors',[])[:2])}, {p.get('year','N/A')}):\n"
        f"{p.get('abstract','')[:500]}"
        for i, p in enumerate(papers)
    )

    comparison = call_granite(
        "Compare and contrast the following research papers. Produce a structured comparison covering: "
        "(1) Research objectives, (2) Methodologies, (3) Datasets used, "
        "(4) Key results, (5) Limitations, (6) Overall strengths and weaknesses. "
        "Conclude with a recommendation on which paper's approach is most rigorous and why.",
        extra_context=papers_block,
    )
    return jsonify({"comparison": comparison})


# ── Research Gap Analysis ─────────────────────────────────────────────────────
@app.route("/api/gaps", methods=["POST"])
def research_gaps():
    data  = request.get_json(force=True)
    topic = data.get("topic", "").strip()
    papers = data.get("papers", [])
    if not topic:
        return jsonify({"error": "Topic required"}), 400

    if not papers and ARXIV_SUPPORT:
        papers = search_arxiv(topic)

    abstracts = "\n\n".join(
        f"- {p.get('title','')} ({p.get('published','')[:4]}): {p.get('abstract','')[:300]}"
        for p in papers[:8]
    )

    analysis = call_granite(
        f"Analyze the current state of research on '{topic}' based on the provided papers. "
        "Identify: (1) Key research gaps not addressed in the literature, "
        "(2) Methodological limitations in existing work, "
        "(3) Under-explored sub-topics or populations, "
        "(4) Contradictions or inconsistencies across studies, "
        "(5) Suggested hypotheses for future investigation, "
        "(6) Recommended research directions with potential impact.",
        extra_context=abstracts,
    )
    return jsonify({"analysis": analysis, "papers": papers})


# ── Document Upload & Analysis ────────────────────────────────────────────────
@app.route("/api/upload", methods=["POST"])
def upload_document():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if not file.filename or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file. Allowed: pdf, docx, txt"}), 400

    filename = secure_filename(file.filename)
    unique_name = f"{uuid.uuid4().hex}_{filename}"
    save_path   = UPLOAD_FOLDER / unique_name
    file.save(str(save_path))

    text = extract_text_from_file(save_path)
    if not text.strip():
        return jsonify({"error": "Could not extract text from the document"}), 422

    action  = request.form.get("action", "summarize")
    prompts = {
        "summarize": (
            "Summarize this research document. Extract: objective, methodology, "
            "datasets, key results, limitations, and contributions."
        ),
        "extract":   (
            "Extract all key information from this paper: "
            "authors (if present), research questions, methodology, datasets, "
            "results, conclusions, and limitations. Format as structured bullet points."
        ),
        "review":    (
            "Based on this document, generate a short critical review: "
            "assess the methodology, validity of conclusions, contribution to the field, "
            "and suggest improvements."
        ),
    }
    prompt  = prompts.get(action, prompts["summarize"])
    result  = call_granite(prompt, extra_context=text)

    # Clean up uploaded file
    try:
        save_path.unlink()
    except Exception:
        pass

    return jsonify({"result": result, "action": action, "filename": filename,
                    "chars_extracted": len(text)})


# ── Citation Manager ──────────────────────────────────────────────────────────
@app.route("/api/citations", methods=["GET"])
def get_citations():
    session_id = request.args.get("session_id", "")
    if not session_id or session_id not in research_sessions:
        return jsonify({"citations": []})
    return jsonify({"citations": research_sessions[session_id]["citations"]})


@app.route("/api/citations/add", methods=["POST"])
def add_citation():
    data = request.get_json(force=True)
    session_id = data.get("session_id", "")
    session = _get_or_create_session(session_id)
    citation = {
        "id":      str(uuid.uuid4()),
        "title":   data.get("title", ""),
        "authors": data.get("authors", []),
        "year":    data.get("year", ""),
        "journal": data.get("journal", ""),
        "url":     data.get("url", ""),
        "doi":     data.get("doi", ""),
        "source":  data.get("source", "manual"),
        "format":  AGENT_INSTRUCTIONS["CITATION_FORMAT"],
        "added":   datetime.now().isoformat(),
    }
    session["citations"].append(citation)
    return jsonify({"citation": citation, "total": len(session["citations"])})


@app.route("/api/citations/delete", methods=["POST"])
def delete_citation():
    data = request.get_json(force=True)
    session_id = data.get("session_id", "")
    cid        = data.get("citation_id", "")
    if session_id in research_sessions:
        research_sessions[session_id]["citations"] = [
            c for c in research_sessions[session_id]["citations"] if c["id"] != cid
        ]
    return jsonify({"ok": True})


@app.route("/api/citations/format", methods=["POST"])
def format_citations():
    data       = request.get_json(force=True)
    citations  = data.get("citations", [])
    fmt        = data.get("format", AGENT_INSTRUCTIONS["CITATION_FORMAT"])
    if not citations:
        return jsonify({"error": "No citations provided"}), 400

    cit_block = "\n".join(
        f"[{i+1}] {', '.join(c.get('authors',[])[:3])} ({c.get('year','')}) "
        f"— '{c.get('title','')}' — {c.get('journal','')}"
        for i, c in enumerate(citations)
    )

    formatted = call_granite(
        f"Format the following citations strictly in {fmt} style. "
        "Return only the formatted reference list, numbered.",
        extra_context=cit_block,
    )
    return jsonify({"formatted": formatted, "format": fmt})


# ── Research Report Generation ────────────────────────────────────────────────
@app.route("/api/report", methods=["POST"])
def generate_report():
    data       = request.get_json(force=True)
    topic      = data.get("topic", "").strip()
    report_type = data.get("type", "full")   # full | executive | abstract
    session_id  = data.get("session_id", str(uuid.uuid4()))
    custom_sections = data.get("sections", [])

    if not topic:
        return jsonify({"error": "Topic required"}), 400

    session   = _get_or_create_session(session_id)
    citations = session.get("citations", [])
    history   = session.get("history", [])

    ctx_parts = []
    if citations:
        ctx_parts.append("Available references:\n" + "\n".join(
            f"- {c.get('title','')} ({c.get('year','')})" for c in citations[:15]
        ))
    if history:
        ctx_parts.append("Previous research discussion:\n" + "\n".join(
            f"Q: {h['user']}\nA: {h['assistant'][:300]}"
            for h in history[-4:]
        ))
    context = "\n\n".join(ctx_parts)

    sections_text = (
        f"Sections: {', '.join(custom_sections)}" if custom_sections
        else "Sections: Abstract, Introduction, Background, Methodology, "
             "Results & Discussion, Conclusion, Future Work, References"
    )

    type_instructions = {
        "full":      "Write a comprehensive, well-structured academic research report.",
        "executive": "Write a concise executive summary (1-2 pages) covering key findings.",
        "abstract":  "Write a structured abstract (250-300 words) for the research topic.",
    }
    instruction = type_instructions.get(report_type, type_instructions["full"])

    report_text = call_granite(
        f"{instruction} Topic: '{topic}'. "
        f"{sections_text}. "
        f"Use {AGENT_INSTRUCTIONS['CITATION_FORMAT']} format for all references. "
        "Cite every factual claim. Include a References section at the end.",
        extra_context=context,
    )

    session["history"].append({
        "user":      f"Generate {report_type} report: {topic}",
        "assistant": report_text,
        "ts":        datetime.now().isoformat(),
        "type":      "report",
    })

    return jsonify({
        "report":     report_text,
        "topic":      topic,
        "type":       report_type,
        "session_id": session_id,
    })


# ── Export Report (PDF / DOCX) ────────────────────────────────────────────────
@app.route("/api/export", methods=["POST"])
def export_report():
    data    = request.get_json(force=True)
    content = data.get("content", "").strip()
    title   = data.get("title", "Research Report")
    fmt     = data.get("format", "pdf").lower()

    if not content:
        return jsonify({"error": "Content required"}), 400

    if fmt == "pdf":
        if not REPORTLAB_SUPPORT:
            return jsonify({"error": "ReportLab not installed"}), 500
        pdf_bytes = _build_pdf_report(title, content)
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
        )
    if fmt == "docx":
        if not DOCX_SUPPORT:
            return jsonify({"error": "python-docx not installed"}), 500
        docx_bytes = _build_docx_report(title, content)
        return send_file(
            io.BytesIO(docx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        )
    return jsonify({"error": "Unsupported format. Use 'pdf' or 'docx'"}), 400


# ── Research History ──────────────────────────────────────────────────────────
@app.route("/api/history", methods=["GET"])
def get_history():
    session_id = request.args.get("session_id", "")
    if not session_id or session_id not in research_sessions:
        return jsonify({"history": [], "title": ""})
    s = research_sessions[session_id]
    return jsonify({"history": s["history"], "title": s["title"],
                    "created": s["created"], "citations": len(s["citations"])})


@app.route("/api/sessions", methods=["GET"])
def list_sessions():
    sessions = [
        {"id": sid, "title": s["title"], "created": s["created"],
         "messages": len(s["history"]), "citations": len(s["citations"])}
        for sid, s in research_sessions.items()
    ]
    sessions.sort(key=lambda x: x["created"], reverse=True)
    return jsonify({"sessions": sessions})


@app.route("/api/sessions/delete", methods=["POST"])
def delete_session():
    data = request.get_json(force=True)
    sid  = data.get("session_id", "")
    if sid in research_sessions:
        del research_sessions[sid]
    return jsonify({"ok": True})


# ─────────────────────────────────────────────────────────────────────────────
# Serve static uploads (if needed)
# ─────────────────────────────────────────────────────────────────────────────
@app.route("/uploads/<path:filename>")
def serve_upload(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    port  = int(os.getenv("PORT", 5000))
    debug = os.getenv("FLASK_DEBUG", "False").lower() in ("1", "true", "yes")
    log.info("Starting AI Research Agent on http://localhost:%d", port)
    log.info("Domain: %s | Citation: %s | Style: %s",
             AGENT_INSTRUCTIONS["RESEARCH_DOMAIN"],
             AGENT_INSTRUCTIONS["CITATION_FORMAT"],
             AGENT_INSTRUCTIONS["WRITING_STYLE"])
    app.run(host="0.0.0.0", port=port, debug=debug)
